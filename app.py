from __future__ import annotations

import os
import sqlite3
import hashlib
import secrets
from datetime import datetime
from pathlib import Path

from docx import Document
from fastapi import FastAPI, File, Form, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse, PlainTextResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from openpyxl import load_workbook
try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except Exception:
    HTML = None
    WEASYPRINT_AVAILABLE = False

BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = Path(os.environ.get("PBL_DATA_DIR", BASE_DIR / "data"))
GENERATED_DIR = DATA_DIR / "generated"
SUBMISSIONS_DIR = DATA_DIR / "submissions"
DEMO_DIR = BASE_DIR / "demo"
DB_PATH = DATA_DIR / "pbl.db"

DATA_DIR.mkdir(exist_ok=True)
GENERATED_DIR.mkdir(exist_ok=True)
SUBMISSIONS_DIR.mkdir(exist_ok=True)

app = FastAPI(title="PBL Manager V0.3.0")
ADMIN_PASSWORD = os.environ.get("PBL_ADMIN_PASSWORD", "pbl123")
ADMIN_COOKIE_TOKEN = hashlib.sha256((ADMIN_PASSWORD + "|pbl-manager-v03").encode()).hexdigest()

@app.middleware("http")
async def protect_admin(request: Request, call_next):
    path = request.url.path
    is_group_admin = path.startswith("/group/") and not path.endswith("/print") and not path.endswith("/pdf")
    protected = path == "/" or path.startswith("/admin") or is_group_admin
    if protected and path not in ("/admin-login", "/admin-logout"):
        if request.cookies.get("pbl_admin") != ADMIN_COOKIE_TOKEN:
            return RedirectResponse(url="/admin-login", status_code=303)
    return await call_next(request)
app.mount("/static", StaticFiles(directory=BASE_DIR / "static"), name="static")
templates = Jinja2Templates(directory=BASE_DIR / "templates")

PROJECT_INFO = {
    1: {"title": "Dự án số 1", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Theo biểu đồ đặc tính tải trọng", "direction": "Làm việc một chiều"},
    2: {"title": "Dự án số 2", "task": "Thiết kế hệ thống cơ khí sử dụng để nâng hạ hàng hóa.", "load": "Tải thay đổi, rung động vừa", "direction": "Làm việc hai chiều"},
    3: {"title": "Dự án số 3", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Tải trọng thay đổi, rung động", "direction": "Làm việc một chiều"},
    4: {"title": "Dự án số 4", "task": "Thiết kế hệ thống cơ khí sử dụng để nâng hạ hàng hóa.", "load": "Tải trọng thay đổi, rung động nhẹ", "direction": "Làm việc hai chiều"},
    5: {"title": "Dự án số 5", "task": "Thiết kế hệ thống cơ khí sử dụng trong vận chuyển hàng hóa.", "load": "Tải trọng thay đổi, rung động nhẹ", "direction": "Làm việc một chiều"},
}

DEFAULT_MILESTONES = [
    (1, "Nhận đề & phân tích nhiệm vụ"),
    (2, "Tính chọn động cơ / phân phối tỉ số truyền"),
    (3, "Thiết kế các bộ truyền và chi tiết chính"),
    (4, "Hoàn thiện bản vẽ & báo cáo"),
]


def db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA foreign_keys = ON")
    return conn


def init_db():
    with db() as c:
        c.executescript(
            """
            CREATE TABLE IF NOT EXISTS students (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                stt INTEGER,
                student_id TEXT UNIQUE NOT NULL,
                name TEXT NOT NULL,
                project_type INTEGER NOT NULL
            );
            CREATE TABLE IF NOT EXISTS project_data (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                project_type INTEGER NOT NULL,
                data_code TEXT NOT NULL,
                p REAL, v REAL, d REAL, t REAL, nam REAL, ngay REAL,
                UNIQUE(project_type, data_code)
            );
            CREATE TABLE IF NOT EXISTS groups_tbl (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_code TEXT UNIQUE NOT NULL,
                project_type INTEGER NOT NULL,
                data_id INTEGER,
                FOREIGN KEY(data_id) REFERENCES project_data(id)
            );
            CREATE TABLE IF NOT EXISTS group_students (
                group_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                PRIMARY KEY(group_id, student_id),
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS templates (
                project_type INTEGER PRIMARY KEY,
                file_path TEXT NOT NULL,
                version INTEGER NOT NULL DEFAULT 1,
                updated_at TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS milestones (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                position INTEGER NOT NULL,
                title TEXT NOT NULL
            );
            CREATE TABLE IF NOT EXISTS group_progress (
                group_id INTEGER NOT NULL,
                milestone_id INTEGER NOT NULL,
                status TEXT NOT NULL DEFAULT 'Chưa bắt đầu',
                score REAL,
                feedback TEXT,
                PRIMARY KEY(group_id, milestone_id),
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(milestone_id) REFERENCES milestones(id) ON DELETE CASCADE
            );
            CREATE TABLE IF NOT EXISTS app_settings (
                key TEXT PRIMARY KEY,
                value TEXT NOT NULL DEFAULT ''
            );
            CREATE TABLE IF NOT EXISTS submissions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                group_id INTEGER NOT NULL,
                student_id INTEGER NOT NULL,
                milestone_id INTEGER,
                note TEXT,
                original_name TEXT NOT NULL,
                stored_name TEXT NOT NULL,
                file_type TEXT NOT NULL,
                uploaded_at TEXT NOT NULL,
                FOREIGN KEY(group_id) REFERENCES groups_tbl(id) ON DELETE CASCADE,
                FOREIGN KEY(student_id) REFERENCES students(id) ON DELETE CASCADE,
                FOREIGN KEY(milestone_id) REFERENCES milestones(id) ON DELETE SET NULL
            );
            """
        )
        if c.execute("SELECT COUNT(*) n FROM milestones").fetchone()["n"] == 0:
            c.executemany("INSERT INTO milestones(position,title) VALUES (?,?)", DEFAULT_MILESTONES)
        defaults = {
            "instructor_name": "TS. Bùi Minh Hiển",
            "start_week": "1",
            "end_week": "15",
            "academic_year_start": "2026",
            "academic_year_end": "2027",
        }
        for key, value in defaults.items():
            c.execute("INSERT OR IGNORE INTO app_settings(key,value) VALUES (?,?)", (key, value))


def parse_students_xlsx(path: Path):
    wb = load_workbook(path, data_only=True)
    ws = wb.active
    header_row = None
    cols = {}
    for r in range(1, min(ws.max_row, 30) + 1):
        norm = [str(ws.cell(r, c).value or "").strip().lower() for c in range(1, ws.max_column + 1)]
        if any(v in ("số thẻ", "mssv", "mã sv", "ma sv") for v in norm) and any(v in ("họ tên", "họ và tên", "ho ten") for v in norm):
            header_row = r
            for i, v in enumerate(norm, 1):
                if v in ("tt", "stt"):
                    cols["stt"] = i
                elif v in ("số thẻ", "mssv", "mã sv", "ma sv"):
                    cols["student_id"] = i
                elif v in ("họ tên", "họ và tên", "ho ten"):
                    cols["name"] = i
                elif v in ("đề", "de", "mã đề", "ma de"):
                    cols["project_type"] = i
            break
    if not header_row or not all(k in cols for k in ("student_id", "name")):
        raise ValueError("Không tìm thấy hàng tiêu đề gồm Số thẻ/MSSV và Họ tên.")
    result = []
    for r in range(header_row + 1, ws.max_row + 1):
        sid = ws.cell(r, cols["student_id"]).value
        name = ws.cell(r, cols["name"]).value
        ptype = ws.cell(r, cols["project_type"]).value if "project_type" in cols else None
        if sid is None and name is None:
            continue
        if sid is None or name is None:
            continue
        if ptype in (None, ""):
            ptype = 0
        else:
            ptype = int(float(ptype))
            if ptype not in PROJECT_INFO:
                raise ValueError(f"Dòng {r}: mã Đề={ptype} chưa được hỗ trợ (chỉ DA1-DA5 hoặc để trống).")
        stt_raw = ws.cell(r, cols["stt"]).value if "stt" in cols else len(result) + 1
        try:
            stt = int(stt_raw)
        except Exception:
            stt = len(result) + 1
        result.append((stt, str(sid).strip(), str(name).strip(), ptype))
    return result


def parse_project_data(path: Path, max_project=5):
    out = []
    if path.suffix.lower() == ".xlsx":
        wb = load_workbook(path, data_only=True)
        for ptype in range(1, max_project + 1):
            name = f"De-{ptype}"
            if name not in wb.sheetnames:
                raise ValueError(f"Thiếu sheet {name}")
            ws = wb[name]
            for r in range(2, ws.max_row + 1):
                vals = [ws.cell(r, c).value for c in range(1, 8)]
                if vals[0] in (None, ""):
                    continue
                out.append((ptype, str(vals[0]).strip(), *vals[1:7]))
    elif path.suffix.lower() == ".xls":
        try:
            import xlrd
        except ImportError as exc:
            raise ValueError("Đọc trực tiếp .xls cần gói xlrd. Có thể Save As file sang .xlsx rồi import lại.") from exc
        wb = xlrd.open_workbook(path)
        for ptype in range(1, max_project + 1):
            name = f"De-{ptype}"
            try:
                ws = wb.sheet_by_name(name)
            except Exception as exc:
                raise ValueError(f"Thiếu sheet {name}") from exc
            for r in range(1, ws.nrows):
                vals = ws.row_values(r, 0, min(7, ws.ncols)) + [None] * max(0, 7 - ws.ncols)
                if vals[0] in (None, ""):
                    continue
                out.append((ptype, str(vals[0]).strip(), *vals[1:7]))
    else:
        raise ValueError("Chỉ hỗ trợ .xls hoặc .xlsx")
    return out


def next_group_code(c):
    rows = c.execute("SELECT group_code FROM groups_tbl").fetchall()
    nums = []
    for r in rows:
        code = str(r["group_code"] or "")
        if code.startswith("N") and code[1:].isdigit():
            nums.append(int(code[1:]))
    return f"N{(max(nums) if nums else 0)+1:02d}"


def assign_next_unused_data(c, group_id: int, project_type: int):
    row = c.execute(
        """SELECT pd.id FROM project_data pd
           WHERE pd.project_type=? AND pd.id NOT IN
           (SELECT data_id FROM groups_tbl WHERE data_id IS NOT NULL AND id<>?)
           ORDER BY pd.id LIMIT 1""",
        (project_type, group_id),
    ).fetchone()
    c.execute("UPDATE groups_tbl SET data_id=? WHERE id=?", (row["id"] if row else None, group_id))


def ensure_progress_rows(c, group_id: int):
    mids = c.execute("SELECT id FROM milestones ORDER BY position").fetchall()
    c.executemany(
        "INSERT OR IGNORE INTO group_progress(group_id,milestone_id) VALUES (?,?)",
        [(group_id, m["id"]) for m in mids],
    )


def regroup_and_assign():
    with db() as c:
        students = c.execute("SELECT * FROM students ORDER BY stt,id").fetchall()
        c.execute("DELETE FROM group_progress")
        c.execute("DELETE FROM group_students")
        c.execute("DELETE FROM groups_tbl")
        group_no = 1
        idx = 0
        while idx < len(students):
            s1 = students[idx]
            if not s1["project_type"]:
                idx += 1
                continue
            members = [s1]
            if idx + 1 < len(students) and students[idx + 1]["project_type"] == s1["project_type"]:
                members.append(students[idx + 1])
                idx += 2
            else:
                idx += 1
            cur = c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,?)", (f"N{group_no:02d}", s1["project_type"]))
            gid = cur.lastrowid
            c.executemany("INSERT INTO group_students(group_id,student_id) VALUES (?,?)", [(gid, m["id"]) for m in members])
            assign_next_unused_data(c, gid, s1["project_type"])
            ensure_progress_rows(c, gid)
            group_no += 1


def import_students(path: Path):
    rows = parse_students_xlsx(path)
    with db() as c:
        c.execute("DELETE FROM group_progress")
        c.execute("DELETE FROM group_students")
        c.execute("DELETE FROM groups_tbl")
        c.execute("DELETE FROM students")
        c.executemany("INSERT INTO students(stt,student_id,name,project_type) VALUES (?,?,?,?)", rows)
    regroup_and_assign()
    return len(rows)


def import_project_data(path: Path):
    rows = parse_project_data(path)
    with db() as c:
        c.execute("UPDATE groups_tbl SET data_id=NULL")
        c.execute("DELETE FROM project_data")
        c.executemany("INSERT INTO project_data(project_type,data_code,p,v,d,t,nam,ngay) VALUES (?,?,?,?,?,?,?,?)", rows)
        # Giữ nguyên phân nhóm thủ công; chỉ cấp lại số liệu theo loại DA hiện tại.
        for ptype in PROJECT_INFO:
            groups = c.execute("SELECT id FROM groups_tbl WHERE project_type=? ORDER BY id", (ptype,)).fetchall()
            for g in groups:
                assign_next_unused_data(c, g["id"], ptype)
                ensure_progress_rows(c, g["id"])
    return len(rows)


def install_template(project_type: int, source_path: Path):
    target = DATA_DIR / f"DA{project_type}_template.docx"
    target.write_bytes(source_path.read_bytes())
    with db() as c:
        row = c.execute("SELECT version FROM templates WHERE project_type=?", (project_type,)).fetchone()
        version = (row["version"] + 1) if row else 1
        c.execute(
            "INSERT INTO templates(project_type,file_path,version,updated_at) VALUES (?,?,?,?) "
            "ON CONFLICT(project_type) DO UPDATE SET file_path=excluded.file_path,version=excluded.version,updated_at=excluded.updated_at",
            (project_type, str(target), version, datetime.now().isoformat(timespec="seconds")),
        )
    return version


def seed_demo_if_empty():
    with db() as c:
        if c.execute("SELECT COUNT(*) n FROM students").fetchone()["n"]:
            return
    for ptype in PROJECT_INFO:
        install_template(ptype, DEMO_DIR / "templates" / f"DA{ptype}_PBL1.docx")
    import_project_data(DEMO_DIR / "dataset.xlsx")
    import_students(DEMO_DIR / "student_list.xlsx")


def get_settings():
    with db() as c:
        rows = c.execute("SELECT key,value FROM app_settings").fetchall()
    return {r["key"]: r["value"] for r in rows}


def update_settings(instructor_name: str, start_week: str, end_week: str, academic_year_start: str, academic_year_end: str):
    values = {
        "instructor_name": instructor_name.strip(),
        "start_week": start_week.strip(),
        "end_week": end_week.strip(),
        "academic_year_start": academic_year_start.strip(),
        "academic_year_end": academic_year_end.strip(),
    }
    for key in ("start_week", "end_week", "academic_year_start", "academic_year_end"):
        if values[key] and not values[key].isdigit():
            raise ValueError(f"{key} phải là số.")
    if values["start_week"] and values["end_week"] and int(values["start_week"]) > int(values["end_week"]):
        raise ValueError("Tuần bắt đầu không được lớn hơn tuần kết thúc.")
    with db() as c:
        for key, value in values.items():
            c.execute("INSERT INTO app_settings(key,value) VALUES (?,?) ON CONFLICT(key) DO UPDATE SET value=excluded.value", (key, value))


def load_image_url(project_type: int):
    return f"/static/project_images/DA{project_type}_load.png"


def system_image_url(project_type: int):
    return f"/static/project_images/DA{project_type}_system.png"


def format_num(v):
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v)


def set_para_text(paragraph, text):
    if paragraph.runs:
        for r in paragraph.runs:
            r.text = ""
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def generate_docx(group_id: int):
    with db() as c:
        g = c.execute(
            """SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay,t.file_path,t.version
               FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id
               LEFT JOIN templates t ON t.project_type=g.project_type WHERE g.id=?""",
            (group_id,),
        ).fetchone()
        if not g:
            raise ValueError("Không tìm thấy nhóm")
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
    if not g["file_path"]:
        raise ValueError("Chưa có template Word cho loại đề này")
    if not g["data_code"]:
        raise ValueError("Nhóm chưa được cấp dòng số liệu")

    doc = Document(g["file_path"])
    names = " - ".join(m["name"] for m in members)
    if len(doc.tables) > 1:
        info = doc.tables[1]
        set_para_text(info.cell(0, 1).paragraphs[0], names)
        set_para_text(info.cell(1, 1).paragraphs[0], f"{g['group_code']} / {g['data_code']}")

    p, v, d, tval, nam, ngay = [format_num(g[k]) for k in ("p", "v", "d", "t", "nam", "ngay")]
    if g["project_type"] in (1, 2, 3, 4) and len(doc.tables) > 2:
        cell = doc.tables[2].cell(1, 0)
        paras = cell.paragraphs
        labels = {
            1: ("Lực kéo băng tải", "Vận tốc băng tải"),
            2: ("Lực kéo dây cáp", "Vận tốc kéo cáp"),
            3: ("Lực kéo băng tải", "Vận tốc băng tải"),
            4: ("Lực kéo cáp", "Vận tốc kéo cáp"),
        }[g["project_type"]]
        lines = [f"1. {labels[0]} :    P = {p}", f"2. {labels[1]} :    V = {v}", f"3. Đường kính tang :    D = {d}"]
        for i in range(min(3, len(paras))):
            set_para_text(paras[i], lines[i])
        for para in paras:
            txt = para.text
            if "Thời gian phục vụ" in txt:
                set_para_text(para, f"5. Thời gian phục vụ: {tval} năm")
            elif "Một năm làm việc" in txt:
                set_para_text(para, f"Một năm làm việc {nam} ngày, một ngày làm việc {ngay} giờ")
    elif g["project_type"] == 5:
        for para in doc.paragraphs:
            txt = para.text
            if txt.startswith("1. Lực kéo băng tải"):
                set_para_text(para, f"1. Lực kéo băng tải :\t\tP = {p}")
            elif txt.startswith("2. Vận tốc băng tải"):
                set_para_text(para, f"2. Vận tốc băng tải :\t\tV = {v}")
            elif txt.startswith("3. Đường kính tang"):
                set_para_text(para, f"3. Đường kính tang :\t\tD = {d}")
            elif txt.startswith("5. Thời gian phục vụ"):
                set_para_text(para, f"5. Thời gian phục vụ: {tval} năm")
            elif txt.startswith("Một năm làm việc"):
                set_para_text(para, f"Một năm làm việc {nam} ngày, một ngày làm việc {ngay} giờ.")

    out = GENERATED_DIR / f"{g['group_code']}_DA{g['project_type']}_{g['data_code'].replace('.', '-')}.docx"
    doc.save(out)
    return out


ALLOWED_SUBMISSION_EXTS = {".docx", ".xlsx", ".xls", ".dwg", ".dxf"}

def safe_filename(name: str):
    import re
    base = Path(name).name
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", base)
    return stem[:160] or "file"

def group_members_text(c, group_id: int):
    rows = c.execute("SELECT s.name FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
    return " - ".join(r["name"] for r in rows)


def render(request: Request, name: str, **kwargs):
    return templates.TemplateResponse(request=request, name=name, context={"PROJECT_INFO": PROJECT_INFO, "format_num": format_num, "settings": get_settings(), "load_image_url": load_image_url, "system_image_url": system_image_url, **kwargs})


def redirect_with(path: str, *, msg: str | None = None, err: str | None = None):
    from urllib.parse import urlencode
    q = urlencode({k: v for k, v in {"msg": msg, "err": err}.items() if v})
    return RedirectResponse(path + ("?" + q if q else ""), status_code=303)


@app.get("/admin-login", response_class=HTMLResponse, name="admin_login")
def admin_login_get(request: Request):
    return render(request, "admin_login.html")

@app.post("/admin-login", name="admin_login_post")
def admin_login_post(password: str = Form(...)):
    if not secrets.compare_digest(password, ADMIN_PASSWORD):
        return redirect_with("/admin-login", err="Mật khẩu giảng viên không đúng.")
    response = RedirectResponse(url="/", status_code=303)
    response.set_cookie("pbl_admin", ADMIN_COOKIE_TOKEN, httponly=True, samesite="lax", max_age=60*60*12)
    return response

@app.get("/admin-logout", name="admin_logout")
def admin_logout():
    response = RedirectResponse(url="/admin-login", status_code=303)
    response.delete_cookie("pbl_admin")
    return response


@app.on_event("startup")
def startup():
    init_db()
    seed_demo_if_empty()


@app.get("/", response_class=HTMLResponse, name="dashboard")
def dashboard(request: Request):
    with db() as c:
        stats = {
            "students": c.execute("SELECT COUNT(*) n FROM students").fetchone()["n"],
            "groups": c.execute("SELECT COUNT(*) n FROM groups_tbl").fetchone()["n"],
            "assigned": c.execute("SELECT COUNT(*) n FROM groups_tbl WHERE data_id IS NOT NULL").fetchone()["n"],
            "templates": c.execute("SELECT COUNT(*) n FROM templates").fetchone()["n"],
        }
        groups = c.execute(
            """SELECT g.*,pd.data_code,
                      (SELECT GROUP_CONCAT(s.name, ' • ') FROM group_students gs JOIN students s ON s.id=gs.student_id WHERE gs.group_id=g.id) members,
                      COALESCE((SELECT ROUND(100.0*SUM(CASE WHEN gp.status='Hoàn thành' THEN 1 ELSE 0 END)/NULLIF(COUNT(*),0)) FROM group_progress gp WHERE gp.group_id=g.id),0) progress
               FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id ORDER BY g.id"""
        ).fetchall()
        unassigned = c.execute("SELECT s.* FROM students s WHERE NOT EXISTS (SELECT 1 FROM group_students gs WHERE gs.student_id=s.id) ORDER BY s.stt,s.id").fetchall()
        all_students = c.execute("SELECT s.*, (SELECT g.group_code FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_code FROM students s ORDER BY s.stt,s.id").fetchall()
    return render(request, "dashboard.html", stats=stats, groups=groups, unassigned=unassigned, all_students=all_students)


@app.get("/admin/import", response_class=HTMLResponse, name="admin_import")
def admin_import_get(request: Request):
    with db() as c:
        trows = c.execute("SELECT * FROM templates ORDER BY project_type").fetchall()
        template_map = {r["project_type"]: r for r in trows}
        counts = c.execute("SELECT project_type,COUNT(*) n FROM project_data GROUP BY project_type ORDER BY project_type").fetchall()
    return render(request, "import.html", template_map=template_map, counts=counts)


@app.post("/admin/import", name="admin_import_post")
async def admin_import_post(kind: str = Form(...), file: UploadFile = File(...), project_type: int | None = Form(None)):
    suffix = Path(file.filename or "").suffix.lower()
    tmp = DATA_DIR / f"upload_{datetime.now().strftime('%Y%m%d%H%M%S%f')}{suffix}"
    try:
        tmp.write_bytes(await file.read())
        if kind == "students":
            if suffix != ".xlsx":
                raise ValueError("Danh sách sinh viên V0.1 yêu cầu file .xlsx")
            n = import_students(tmp)
            return redirect_with("/admin/import", msg=f"Đã import {n} sinh viên và tự động chia nhóm.")
        if kind == "dataset":
            n = import_project_data(tmp)
            return redirect_with("/admin/import", msg=f"Đã import {n} dòng số liệu cho DA1-DA5.")
        if kind == "template":
            if project_type not in PROJECT_INFO or suffix != ".docx":
                raise ValueError("Template phải là .docx và DA từ 1 đến 5.")
            version = install_template(int(project_type), tmp)
            return redirect_with("/admin/import", msg=f"Đã cập nhật template DA{project_type}, phiên bản {version}.")
        raise ValueError("Loại import không hợp lệ")
    except Exception as exc:
        return redirect_with("/admin/import", err=str(exc))
    finally:
        if tmp.exists():
            tmp.unlink()


@app.post("/admin/settings", name="admin_settings_post")
def admin_settings_post(
    instructor_name: str = Form(""),
    start_week: str = Form(""),
    end_week: str = Form(""),
    academic_year_start: str = Form(""),
    academic_year_end: str = Form(""),
):
    try:
        update_settings(instructor_name, start_week, end_week, academic_year_start, academic_year_end)
    except Exception as exc:
        return redirect_with("/admin/import", err=str(exc))
    return redirect_with("/admin/import", msg="Đã lưu thông tin học phần và giảng viên hướng dẫn.")


@app.get("/admin/groups", response_class=HTMLResponse, name="admin_groups")
def admin_groups(request: Request):
    with db() as c:
        groups = c.execute("""SELECT g.*,pd.data_code,
            (SELECT GROUP_CONCAT(s.name, ' • ') FROM group_students gs JOIN students s ON s.id=gs.student_id WHERE gs.group_id=g.id) members,
            (SELECT COUNT(*) FROM group_students gs WHERE gs.group_id=g.id) member_count
            FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id ORDER BY g.group_code""").fetchall()
        students = c.execute("""SELECT s.*, (SELECT g.id FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_id,
            (SELECT g.group_code FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=s.id LIMIT 1) group_code
            FROM students s ORDER BY s.stt,s.id""").fetchall()
    return render(request, "admin_groups.html", groups=groups, students=students)

@app.post("/admin/groups/move", name="admin_group_move")
def admin_group_move(student_db_id: int = Form(...), target_group_id: int = Form(0), new_project_type: int = Form(0)):
    try:
        with db() as c:
            student = c.execute("SELECT * FROM students WHERE id=?", (student_db_id,)).fetchone()
            if not student:
                raise ValueError("Không tìm thấy sinh viên.")
            old = c.execute("SELECT g.id FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=?", (student_db_id,)).fetchone()
            if target_group_id:
                target = c.execute("SELECT * FROM groups_tbl WHERE id=?", (target_group_id,)).fetchone()
                if not target:
                    raise ValueError("Không tìm thấy nhóm đích.")
                count = c.execute("SELECT COUNT(*) n FROM group_students WHERE group_id=?", (target_group_id,)).fetchone()["n"]
                if count >= 2 and (not old or old["id"] != target_group_id):
                    raise ValueError("Nhóm đích đã có 2 sinh viên.")
                if old and old["id"] == target_group_id:
                    return redirect_with("/admin/groups", msg="Sinh viên đã ở nhóm này.")
                if old:
                    c.execute("DELETE FROM group_students WHERE group_id=? AND student_id=?", (old["id"], student_db_id))
                c.execute("INSERT OR IGNORE INTO group_students(group_id,student_id) VALUES (?,?)", (target_group_id, student_db_id))
                c.execute("UPDATE students SET project_type=? WHERE id=?", (target["project_type"], student_db_id))
            else:
                if new_project_type not in PROJECT_INFO:
                    raise ValueError("Khi tạo nhóm mới cần chọn DA1-DA5.")
                if old:
                    c.execute("DELETE FROM group_students WHERE group_id=? AND student_id=?", (old["id"], student_db_id))
                code = next_group_code(c)
                cur = c.execute("INSERT INTO groups_tbl(group_code,project_type) VALUES (?,?)", (code, new_project_type))
                target_group_id = cur.lastrowid
                c.execute("INSERT INTO group_students(group_id,student_id) VALUES (?,?)", (target_group_id, student_db_id))
                c.execute("UPDATE students SET project_type=? WHERE id=?", (new_project_type, student_db_id))
                assign_next_unused_data(c, target_group_id, new_project_type)
                ensure_progress_rows(c, target_group_id)
            if old:
                nleft = c.execute("SELECT COUNT(*) n FROM group_students WHERE group_id=?", (old["id"],)).fetchone()["n"]
                if nleft == 0:
                    c.execute("DELETE FROM groups_tbl WHERE id=?", (old["id"],))
            ensure_progress_rows(c, target_group_id)
        return redirect_with("/admin/groups", msg="Đã cập nhật phân công sinh viên/nhóm.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))

@app.post("/admin/groups/{group_id}/project", name="admin_group_project")
def admin_group_project(group_id: int, project_type: int = Form(...)):
    try:
        if project_type not in PROJECT_INFO:
            raise ValueError("Loại đề phải từ DA1 đến DA5.")
        with db() as c:
            g = c.execute("SELECT * FROM groups_tbl WHERE id=?", (group_id,)).fetchone()
            if not g:
                raise ValueError("Không tìm thấy nhóm.")
            c.execute("UPDATE groups_tbl SET project_type=?, data_id=NULL WHERE id=?", (project_type, group_id))
            c.execute("UPDATE students SET project_type=? WHERE id IN (SELECT student_id FROM group_students WHERE group_id=?)", (project_type, group_id))
            assign_next_unused_data(c, group_id, project_type)
        return redirect_with("/admin/groups", msg="Đã đổi loại đề và cấp lại bộ số liệu phù hợp.")
    except Exception as exc:
        return redirect_with("/admin/groups", err=str(exc))


@app.get("/group/{group_id}", response_class=HTMLResponse, name="group_detail")
def group_detail(request: Request, group_id: int):
    with db() as c:
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id WHERE g.id=?", (group_id,)).fetchone()
        if not g:
            return PlainTextResponse("Không tìm thấy nhóm", status_code=404)
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (group_id,)).fetchall()
        submissions = c.execute("""SELECT sub.*,s.name student_name,m.title milestone_title FROM submissions sub
            JOIN students s ON s.id=sub.student_id LEFT JOIN milestones m ON m.id=sub.milestone_id
            WHERE sub.group_id=? ORDER BY sub.uploaded_at DESC""", (group_id,)).fetchall()
    return render(request, "group.html", g=g, members=members, progress=progress, submissions=submissions, info=PROJECT_INFO[g["project_type"]])


@app.post("/group/{group_id}/progress/{milestone_id}", name="update_progress")
def update_progress(group_id: int, milestone_id: int, status: str = Form(...), score: str = Form(""), feedback: str = Form("")):
    try:
        score_value = float(score) if score.strip() else None
    except ValueError:
        return redirect_with(f"/group/{group_id}", err="Điểm phải là số.")
    with db() as c:
        c.execute("INSERT INTO group_progress(group_id,milestone_id,status,score,feedback) VALUES (?,?,?,?,?) ON CONFLICT(group_id,milestone_id) DO UPDATE SET status=excluded.status,score=excluded.score,feedback=excluded.feedback", (group_id, milestone_id, status, score_value, feedback.strip()))
    return redirect_with(f"/group/{group_id}", msg="Đã cập nhật tiến độ.")


@app.get("/student", response_class=HTMLResponse, name="student_lookup")
def student_lookup_get(request: Request):
    return render(request, "student_lookup.html")


@app.post("/student", name="student_lookup_post")
def student_lookup_post(student_id: str = Form(...)):
    sid = student_id.strip()
    with db() as c:
        row = c.execute("SELECT s.student_id FROM students s WHERE s.student_id=?", (sid,)).fetchone()
    if not row:
        return redirect_with("/student", err="Không tìm thấy số thẻ sinh viên.")
    return RedirectResponse(f"/student/{sid}", status_code=303)


@app.get("/student/{student_id}", response_class=HTMLResponse, name="student_portal")
def student_portal(request: Request, student_id: str):
    with db() as c:
        s = c.execute("SELECT * FROM students WHERE student_id=?", (student_id,)).fetchone()
        if not s:
            return PlainTextResponse("Không tìm thấy sinh viên", status_code=404)
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id LEFT JOIN project_data pd ON pd.id=g.data_id WHERE gs.student_id=?", (s["id"],)).fetchone()
        if not g:
            return render(request, "student_unassigned.html", s=s)
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (g["id"],)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (g["id"],)).fetchall()
        submissions = c.execute("""SELECT sub.*,s.name student_name,m.title milestone_title FROM submissions sub
            JOIN students s ON s.id=sub.student_id LEFT JOIN milestones m ON m.id=sub.milestone_id
            WHERE sub.group_id=? ORDER BY sub.uploaded_at DESC""", (g["id"],)).fetchall()
    done = sum(1 for x in progress if x["status"] == "Hoàn thành")
    percent = round(100 * done / len(progress)) if progress else 0
    return render(request, "student_portal.html", s=s, g=g, members=members, progress=progress, submissions=submissions, percent=percent, info=PROJECT_INFO[g["project_type"]])


@app.post("/student/{student_id}/submit", name="student_submit")
async def student_submit(student_id: str, milestone_id: int = Form(...), note: str = Form(""), files: list[UploadFile] = File(...)):
    try:
        with db() as c:
            srow = c.execute("SELECT * FROM students WHERE student_id=?", (student_id,)).fetchone()
            if not srow:
                raise ValueError("Không tìm thấy sinh viên.")
            grow = c.execute("SELECT g.* FROM groups_tbl g JOIN group_students gs ON gs.group_id=g.id WHERE gs.student_id=?", (srow["id"],)).fetchone()
            if not grow:
                raise ValueError("Sinh viên chưa được phân nhóm.")
            valid_mid = c.execute("SELECT id FROM milestones WHERE id=?", (milestone_id,)).fetchone()
            if not valid_mid:
                raise ValueError("Mốc tiến độ không hợp lệ.")
            saved = 0
            for upload in files:
                original = upload.filename or ""
                ext = Path(original).suffix.lower()
                if ext not in ALLOWED_SUBMISSION_EXTS:
                    raise ValueError(f"Không hỗ trợ file {original}. Chỉ nhận DOCX, XLS/XLSX, DWG/DXF.")
                content = await upload.read()
                if len(content) > 50 * 1024 * 1024:
                    raise ValueError(f"File {original} vượt quá 50 MB.")
                stamp = datetime.now().strftime("%Y%m%d_%H%M%S_%f")
                stored = f"G{grow['id']}_S{srow['id']}_{stamp}_{safe_filename(original)}"
                (SUBMISSIONS_DIR / stored).write_bytes(content)
                c.execute("""INSERT INTO submissions(group_id,student_id,milestone_id,note,original_name,stored_name,file_type,uploaded_at)
                    VALUES (?,?,?,?,?,?,?,?)""", (grow["id"], srow["id"], milestone_id, note.strip(), original, stored, ext, datetime.now().isoformat(timespec="seconds")))
                saved += 1
            c.execute("""INSERT INTO group_progress(group_id,milestone_id,status) VALUES (?,?,?)
                ON CONFLICT(group_id,milestone_id) DO UPDATE SET status=CASE WHEN group_progress.status='Hoàn thành' THEN group_progress.status ELSE excluded.status END""",
                (grow["id"], milestone_id, "Đang thực hiện"))
        return redirect_with(f"/student/{student_id}", msg=f"Đã gửi báo cáo tiến độ và tải lên {saved} file.")
    except Exception as exc:
        return redirect_with(f"/student/{student_id}", err=str(exc))

@app.get("/submission/{submission_id}/download", name="download_submission")
def download_submission(submission_id: int):
    with db() as c:
        row = c.execute("SELECT * FROM submissions WHERE id=?", (submission_id,)).fetchone()
    if not row:
        return PlainTextResponse("Không tìm thấy file", status_code=404)
    path = SUBMISSIONS_DIR / row["stored_name"]
    if not path.exists():
        return PlainTextResponse("File không còn trên máy chủ", status_code=404)
    return FileResponse(path, filename=row["original_name"])


def assignment_context(group_id: int):
    with db() as c:
        g = c.execute("SELECT g.*,pd.data_code,pd.p,pd.v,pd.d,pd.t,pd.nam,pd.ngay FROM groups_tbl g LEFT JOIN project_data pd ON pd.id=g.data_id WHERE g.id=?", (group_id,)).fetchone()
        if not g:
            raise ValueError("Không tìm thấy nhóm")
        members = c.execute("SELECT s.* FROM students s JOIN group_students gs ON gs.student_id=s.id WHERE gs.group_id=? ORDER BY s.stt", (group_id,)).fetchall()
        progress = c.execute("SELECT m.*,gp.status,gp.score,gp.feedback FROM milestones m LEFT JOIN group_progress gp ON gp.milestone_id=m.id AND gp.group_id=? ORDER BY m.position", (group_id,)).fetchall()
    return {"g": g, "members": members, "progress": progress, "info": PROJECT_INFO[g["project_type"]], "settings": get_settings()}


@app.get("/group/{group_id}/print", response_class=HTMLResponse, name="print_assignment")
def print_assignment(request: Request, group_id: int):
    try:
        ctx = assignment_context(group_id)
        return render(request, "assignment_print.html", **ctx)
    except Exception as exc:
        return redirect_with(f"/group/{group_id}", err=str(exc))


@app.get("/group/{group_id}/pdf", name="download_pdf")
def download_pdf(request: Request, group_id: int):
    # WeasyPrint needs native GTK/Pango libraries on Windows. If they are not
    # available, use the browser's built-in Print -> Save as PDF instead.
    if not WEASYPRINT_AVAILABLE:
        return RedirectResponse(url=f"/group/{group_id}/print", status_code=303)
    try:
        ctx = assignment_context(group_id)
        html = templates.get_template("assignment_pdf.html").render(
            request=request, PROJECT_INFO=PROJECT_INFO, format_num=format_num,
            load_image_url=load_image_url, system_image_url=system_image_url, base_dir=str(BASE_DIR), **ctx
        )
        out = GENERATED_DIR / f"{ctx['g']['group_code']}_DA{ctx['g']['project_type']}_{str(ctx['g']['data_code']).replace('.', '-')}.pdf"
        HTML(string=html, base_url=str(BASE_DIR)).write_pdf(out)
    except Exception as exc:
        return redirect_with(f"/group/{group_id}", err=str(exc))
    return FileResponse(out, media_type="application/pdf", filename=out.name)


if __name__ == "__main__":
    init_db()
    seed_demo_if_empty()
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", "5000")))
